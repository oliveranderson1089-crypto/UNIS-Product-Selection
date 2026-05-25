"""Word .docx extractor — paragraphs + tables via python-docx."""

from __future__ import annotations

from pathlib import Path

from .base import ExtractedContent, ExtractedTable


def extract_docx(path: Path) -> ExtractedContent:
    """
    Extract text and tables from a .docx file.

    Note: legacy .doc (Word 97–2003) is NOT supported by python-docx — convert
    to .docx first (LibreOffice headless: `soffice --convert-to docx ...`).
    """
    from docx import Document

    doc = Document(str(path))

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    tables: list[ExtractedTable] = []
    for tbl in doc.tables:
        rows = [
            [cell.text.strip().replace("\n", " ") for cell in row.cells]
            for row in tbl.rows
        ]
        if len(rows) >= 2:
            tables.append(ExtractedTable(rows=rows))

    return ExtractedContent(
        source=path,
        kind="docx",
        text="\n".join(paragraphs),
        tables=tables,
    )


__all__ = ["extract_docx"]
